import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{border_name}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)
        borders.append(element)

def format_doc(docx_file):
    doc = Document(docx_file)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        if len(table.columns) <= 3:
            continue
        for row in table.rows[1:]:  # Skip the header row
            if not row.cells:
                continue
            last_cell = row.cells[-1]
            if not last_cell.text.strip():  # Skip empty cells
                continue

            last_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in last_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER    
        # col_count = len(table.columns)
        # for row in table.rows:
        #     last_cell = row.cells[col_count - 1]
        #     last_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     for para in last_cell.paragraphs:
        #         para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #         for run in para.runs:
        #             run.font.size = Pt(12)
            set_cell_border(last_cell)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("SDG Word Table Formatter (Last Column Border + Center Alignment)")

uploaded_file = st.file_uploader("Upload a Word document (.docx)", type="docx")
if uploaded_file:
    formatted_file = format_doc(uploaded_file)
    st.download_button(
        label="Download Formatted Word File",
        data=formatted_file,
        file_name="Formatted_" + uploaded_file.name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
